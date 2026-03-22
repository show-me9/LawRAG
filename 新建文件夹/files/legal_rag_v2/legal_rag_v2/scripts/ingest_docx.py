"""
Word 文档（.docx）直接摄入脚本。
无需 MinerU 转换，直接解析 Word 文档按条文切分后写入 ChromaDB。

支持文档结构：
    第一章　总则
    第一条　为了...（条文内容）
    第二条　...

用法：
    python scripts/ingest_docx.py
    python scripts/ingest_docx.py --dir ./data/raw_docx   # 指定目录
    python scripts/ingest_docx.py --file ./data/raw_docx/危险化学品安全法.docx  # 单文件
"""
import os
import sys
import re
import time
import argparse
import glob

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import chromadb
from docx import Document as DocxDocument
from tqdm import tqdm
from langchain_core.documents import Document
from langchain_chroma import Chroma

from rag.embeddings import get_embedding_model
from rag.retriever import save_bm25_index
from config import CHROMA_PERSIST_DIR, CHROMA_COLLECTION

# ── 正则：识别章节标题和条文编号 ──────────────────────────
# 匹配：第一章、第二章、第十章 ... 等章节标题
RE_CHAPTER = re.compile(r'^第[一二三四五六七八九十百]+章\s*[\u4e00-\u9fff　\s]+')

# 匹配：第一条、第二条 ... 第一百二十三条 等条文开头
RE_ARTICLE = re.compile(r'^(第[一二三四五六七八九十百零千]+条)\s*(.+)')

# 匹配目录行（如"第一章　总则"出现在文档开头目录区域，需跳过）
RE_TOC_LINE = re.compile(r'^第[一二三四五六七八九十百]+章\s*[\u4e00-\u9fff　]+$')


def extract_law_name_from_docx(doc: DocxDocument) -> str:
    """
    从 Word 文档中提取法规名称。
    通常法规名称是文档中第一个非空段落（正文最大字号的标题）。
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 4:
            # 跳过纯数字、纯符号段落
            if re.search(r'[\u4e00-\u9fff]', text):
                return text
    return "未知法规"


def is_toc_section(paragraphs: list, idx: int, window: int = 30) -> bool:
    """
    判断当前段落是否处于目录区域。
    目录区域特征：连续出现多个"第X章"或"第X条"的短行，且不带正文内容。
    """
    # 简单策略：文档前 window 段中，若出现"目　　录"则认为接下来是目录
    start = max(0, idx - window)
    for p in paragraphs[start:idx]:
        if '目' in p and '录' in p and len(p.strip()) <= 10:
            return True
    return False


def parse_docx(docx_path: str) -> list:
    """
    解析单个 Word 文档，按"条"切分为 Document 列表。

    返回的每个 Document：
        page_content = 条文全文
        metadata = {
            law_name:    法规名称
            chapter:     所属章节
            article:     条文编号（如"第一条"）
            citation:    完整引用字符串
            source_file: 文件名
        }
    """
    doc = DocxDocument(docx_path)
    law_name    = extract_law_name_from_docx(doc)
    source_file = os.path.basename(docx_path)

    # 提取所有段落文本，过滤空行
    paragraphs = [p.text for p in doc.paragraphs]

    # ── 第一遍：跳过目录区域，找到正文起始位置 ──────────────
    body_start = 0
    for i, text in enumerate(paragraphs):
        stripped = text.strip()
        # 目录结束的标志：出现"第一章"且其后跟随正文内容（不是纯目录行）
        if RE_CHAPTER.match(stripped) and i > 5:
            # 往后看几行，若有条文内容则认为正文开始
            lookahead = " ".join(p.strip() for p in paragraphs[i:i+5])
            if RE_ARTICLE.search(lookahead):
                body_start = i
                break

    # ── 第二遍：按章节和条文切分 ─────────────────────────────
    docs       = []
    cur_chapter = ""
    cur_article = ""
    cur_lines   = []   # 当前条文的内容行

    def flush_article():
        """将当前积累的条文内容保存为一个 Document。"""
        if not cur_article or not cur_lines:
            return
        content = cur_article + "　" + "".join(cur_lines).strip()
        if len(content.strip()) < 15:
            return
        citation = " · ".join(filter(None, [law_name, cur_chapter, cur_article]))
        docs.append(Document(
            page_content=content,
            metadata={
                "law_name":    law_name,
                "chapter":     cur_chapter,
                "article":     cur_article,
                "citation":    citation,
                "source_file": source_file,
            }
        ))

    for text in paragraphs[body_start:]:
        stripped = text.strip()
        if not stripped:
            continue

        # 识别章节标题
        if RE_CHAPTER.match(stripped):
            flush_article()
            cur_chapter = stripped
            cur_article = ""
            cur_lines   = []
            continue

        # 识别条文开头
        m = RE_ARTICLE.match(stripped)
        if m:
            flush_article()
            cur_article = m.group(1)   # 例："第一条"
            cur_lines   = [m.group(2)] # 条文第一行内容
            continue

        # 条文续行（同一条文跨多个段落）
        if cur_article:
            cur_lines.append(stripped)

    # 处理最后一条
    flush_article()

    return docs


def ingest_docx_files(paths: list) -> list:
    """解析多个 Word 文档，返回所有 Document。"""
    all_docs = []
    for path in tqdm(paths, desc="解析 Word 文档"):
        try:
            docs = parse_docx(path)
            all_docs.extend(docs)
            print(f"  {os.path.basename(path)}: 切分出 {len(docs)} 个条文")
        except Exception as e:
            print(f"  [警告] 解析失败 {os.path.basename(path)}: {e}")
    return all_docs


def write_to_chroma(docs: list) -> None:
    """将文档批量写入 ChromaDB。"""
    os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
    client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
    embedding_model = get_embedding_model()

    vectorstore = Chroma(
        client=client,
        collection_name=CHROMA_COLLECTION,
        embedding_function=embedding_model,
    )

    batch_size = 100
    total = len(docs)
    print(f"\n开始向量化写入（共 {total} 个条文，每批 {batch_size} 个）...")

    for i in range(0, total, batch_size):
        batch = docs[i: i + batch_size]
        vectorstore.add_documents(batch)
        print(f"  进度：{min(i + batch_size, total)}/{total}")
        time.sleep(0.5)  # 避免触发 API 限流

    # 验证
    col = client.get_collection(CHROMA_COLLECTION)
    print(f"\n写入完成！ChromaDB 中现共有 {col.count()} 条向量记录")


# ── 主流程 ───────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 文档摄入脚本")
    parser.add_argument("--dir",  default="./data/raw_docx", help="Word 文档目录")
    parser.add_argument("--file", default="",                help="单个 Word 文档路径")
    args = parser.parse_args()

    print("=" * 50)
    print("  Word 文档摄入脚本")
    print("=" * 50)

    # 收集要处理的文件列表
    if args.file:
        if not os.path.exists(args.file):
            print(f"[错误] 文件不存在：{args.file}")
            sys.exit(1)
        docx_files = [args.file]
    else:
        docx_files = glob.glob(os.path.join(args.dir, "**", "*.docx"), recursive=True)
        if not docx_files:
            print(f"[错误] 未在 {args.dir} 中找到 .docx 文件")
            sys.exit(1)

    print(f"共发现 {len(docx_files)} 个 Word 文档\n")

    # 解析切分
    all_docs = ingest_docx_files(docx_files)
    if not all_docs:
        print("[错误] 未解析出任何条文，请检查文档格式")
        sys.exit(1)

    print(f"\n共解析出 {len(all_docs)} 个条文")

    # 写入 ChromaDB
    write_to_chroma(all_docs)
    save_bm25_index(all_docs)   # 同步保存 BM25 索引到磁盘
    print("\n摄入完成！可以运行 main.py 开始问答。")
